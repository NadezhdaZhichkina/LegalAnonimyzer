import re
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import streamlit as st
from io import BytesIO

# === Маскировка ===
def remove_sensitive_data(text):
    text = text.replace('\t', ' ').replace('\n', ' ')

    patterns = [
        r'/[А-ЯЁ][а-яё]+ [А-ЯЁ]\.[А-ЯЁ]\./',  # /Фамилия И.О./
        r'\b[А-ЯЁ][а-яё]+[\s]+[А-ЯЁ][а-яё]+[\s]+[А-ЯЁ][а-яё]+\b',  # ФИО с пробелами
        r'\b[А-ЯЁ][а-яё]+ [А-ЯЁ]\.[А-ЯЁ]\.\b',  # Фамилия И.О.
        r'\b[А-ЯЁ]\.[А-ЯЁ]\. ?[А-ЯЁ][а-яё]+\b',  # И.О. Фамилия
        r'\bИНН ?\d{10,12}\b',
        r'\bОГРН ?\d{13}\b',
        r'\bКПП ?\d{9}\b',
        r'\b\d{2} \d{2} \d{6}\b',
        r'\b\d{3}-\d{3}-\d{3} \d{2}\b',
        r'\b\d{2}\.\d{2}\.\d{4}\b',
        r'\b[\w\.-]+@[\w\.-]+\.\w+\b',
        r'(?:(?:8|\+7)[\s\-]?)?\(?\d{3,4}\)?[\s\-]?\d{2,3}[\s\-]?\d{2}[\s\-]?\d{2,4}',
        r'https?://[^\s,]+',
        r'\b[\w\.-]+\.(ru|com|pdf|org|net)\b',
        r'(?i)(юридический|почтовый)?\s*адрес\s*[:\-\u2013\u2014]?\s*[^\n\r]*',
        r'(?i)местонахождение\s*[:\-\u2013\u2014]?\s*[^\n\r]*',
        r'\b\d{6},?\s*город\s+[А-ЯЁа-яё\s]+,?\s*[^,\n\r]{0,100}д[ом]*\.?\s*\d+[^,\n\r]*',
        r'(?i)(расчетный счет|р/с|к/с|кор/с|корреспондентский счет)[^\n\r]*',
        r'\b(ООО|АО|ПАО|ОАО|ЗАО|ИП|ТСЖ|МУП|ГУП|НКО|МКА|СНТ|ПК|КФХ)\s*[«„“"‟”][^«»„“"‟”]*[»”"]',
        r'\b(ПУБЛИЧНОЕ\s+|НЕПУБЛИЧНОЕ\s+)?АКЦИОНЕРНОЕ\s+ОБЩЕСТВО\s+[«"„“‟”][^»"“”]+[»"“”][^,)]*',
        r'\b(ПУБЛИЧНОЕ\s+АКЦИОНЕРНОЕ\s+ОБЩЕСТВО|АКЦИОНЕРНОЕ\s+ОБЩЕСТВО|ОБЩЕСТВО\s+С\s+ОГРАНИЧЕННОЙ\s+ОТВЕТСТВЕННОСТЬЮ)\s+["«„“‟”][^"»“”]+["»“”]',
        r'\bг\.?\s*Москва\b[^\n\r]{0,100}д\.?\s*\d+[^\n\r]*',
        r'\bМосква\b[^\n\r]{0,100}д\.?\s*\d+[^\n\r]*',
        r'(?i)\bнепубличное\s+акционерное\s+общество\s+[«"„“‟”][^«»"“”]+[»"“”]',
        r'(?i)\b(публичное\s+|непубличное\s+)?акционерное\s+общество\s+«(?:[^«»]*«[^«»]*»[^«»]*)+»[^\n\r,)]*',
        r'\b(ПУБЛИЧНОЕ\s+|НЕПУБЛИЧНОЕ\s+)?АКЦИОНЕРНОЕ\s+ОБЩЕСТВО\s+[«"„“‟”][^»"“”\n\r]+(?:[«"„“‟”][^»"“”\n\r]+[»"“”])?[^,;\n\r]*',
        r'\b(АО|ПАО|ОАО|ЗАО)\s+(СК|БАНК|КОМПАНИЯ)?\s*[«"„“‟”][^»"“”\n\r]+(?:[«"„“‟”][^»"“”\n\r]+[»"“”])?[^,;\n\r]*',


    ]

    replacements = {}
    for i, pattern in enumerate(patterns):
        def repl(match):
            key = f"[УДАЛЕНО_{i}_{len(replacements)}]"
            replacements[key] = match.group()
            return key
        text = re.sub(pattern, repl, text)
    return text, replacements

def restore_sensitive_data(text, replacements):
    for key, original in replacements.items():
        text = text.replace(key, original)
    return text

def set_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(11)

# === DOCX обработка ===
def sanitize_docx(file_bytes):
    file_bytes.seek(0)
    doc = Document(BytesIO(file_bytes.read()))
    full_replacements = {}
    for para in doc.paragraphs:
        new_text, replacements = remove_sensitive_data(para.text)
        para.text = new_text
        set_font(para)
        full_replacements.update(replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                new_text, replacements = remove_sensitive_data(cell.text)
                cell.text = new_text
                for p in cell.paragraphs:
                    set_font(p)
                full_replacements.update(replacements)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, full_replacements

def restore_docx(file_bytes, replacements):
    file_bytes.seek(0)
    doc = Document(BytesIO(file_bytes.read()))
    for para in doc.paragraphs:
        para.text = restore_sensitive_data(para.text, replacements)
        set_font(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = restore_sensitive_data(cell.text, replacements)
                for p in cell.paragraphs:
                    set_font(p)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === PDF → DOCX через pdfplumber ===
def convert_pdfplumber_to_docx(file_bytes):
    file_bytes.seek(0)
    doc = Document()
    full_replacements = {}

    with pdfplumber.open(file_bytes) as pdf:
        for page in pdf.pages:
            # Текст
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    clean_text, replacements = remove_sensitive_data(line)
                    para = doc.add_paragraph(clean_text)
                    set_font(para)
                    full_replacements.update(replacements)

            # Таблицы
            for table in page.extract_tables():
                doc_table = doc.add_table(rows=0, cols=len(table[0]))
                for row_data in table:
                    row = doc_table.add_row().cells
                    for i, cell in enumerate(row_data):
                        clean_text, replacements = remove_sensitive_data(cell or "")
                        row[i].text = clean_text
                        full_replacements.update(replacements)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, full_replacements

# === Карта замен ===
def export_replacements(replacements):
    rep_text = "\n".join(f"{k} → {v}" for k, v in replacements.items())
    buffer = BytesIO()
    buffer.write(rep_text.encode("utf-8"))
    buffer.seek(0)
    return buffer

def import_replacements(text_file):
    content = text_file.read().decode("utf-8")
    replacements = {}
    for line in content.strip().splitlines():
        if "→" in line:
            key, val = line.split("→", 1)
            replacements[key.strip()] = val.strip()
    return replacements

# === Интерфейс Streamlit ===
st.set_page_config(page_title="AsiwiAnonymizer", page_icon="📄")
st.title("AsiwiAnonymizer — Обезличивание и восстановление .docx и .pdf")

tab1, tab2 = st.tabs(["🔒 Обезличить", "♻️ Восстановить"])

with tab1:
    uploaded_file = st.file_uploader("Загрузите файл .docx или .pdf", type=["docx", "pdf"])
    if uploaded_file and st.button("Обезличить"):
        filetype = uploaded_file.name.split(".")[-1].lower()
        if filetype == "docx":
            output, replacements = sanitize_docx(uploaded_file)
            filename = "обезличенный.docx"
        elif filetype == "pdf":
            output, replacements = convert_pdfplumber_to_docx(uploaded_file)
            filename = "обезличенный.docx"
        else:
            st.error("Неподдерживаемый формат.")
            st.stop()

        st.download_button("📁 Скачать обезличенный файл", data=output, file_name=filename)
        if replacements:
            rep_file = export_replacements(replacements)
            st.download_button("🗂️ Скачать карту замен (.txt)", data=rep_file, file_name="карта_замен.txt")

with tab2:
    file_docx = st.file_uploader("Загрузите изменённый обезличенный .docx", type=["docx"], key="restore_docx")
    file_map = st.file_uploader("Загрузите соответствующую карту замен (.txt)", type=["txt"], key="restore_map")
    if file_docx and file_map and st.button("Восстановить"):
        replacements = import_replacements(file_map)
        restored_doc = restore_docx(file_docx, replacements)
        st.download_button("📄 Скачать восстановленный .docx", data=restored_doc, file_name="восстановленный.docx")
